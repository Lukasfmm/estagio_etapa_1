"""
============================================================================
Módulo: build_report.py
Projeto: Cockpit de Relatórios
Autor: Lucas Ferreira Mendes Moraes
Data: Outubro 2025
============================================================================

DESCRIÇÃO:
    Módulo de serviço responsável por gerar relatórios em DOCX e PDF.
    Carrega dados dos CSVs gerados pelo ETL, calcula métricas e percentuais,
    preenche o template DOCX e converte para PDF.

FUNCIONALIDADES:
    - Carregamento inteligente de todos os CSVs
    - Cálculo de métricas derivadas e percentuais
    - Seleção dinâmica de dados baseada no tipo de relatório
    - Preenchimento automático do template DOCX
    - Conversão para PDF
    - População de tabelas com formatação profissional

TIPOS DE RELATÓRIO SUPORTADOS:
    - Nacional (visão consolidada)
    - Regional (por região geográfica)
    - Por Setor (por setor comercial)
    - Por Grupo (por grupo empresarial)
    - Por Marca (por marca de veículo)
    - Por PDV (por concessionária específica)

DEPENDÊNCIAS:
    - pandas: Manipulação de dados
    - python-docx: Geração de documentos Word
    - docx2pdf: Conversão para PDF
============================================================================
"""

# ============================================================================
# IMPORTAÇÕES
# ============================================================================
import pandas as pd
from pathlib import Path
import docx
import docx2pdf
from datetime import datetime
import locale
import logging

# ============================================================================
# CONFIGURAÇÕES E CAMINHOS GLOBAIS
# ============================================================================
BASE_DIR = Path(__file__).resolve().parent              # pasta report/
DATA_DIR = BASE_DIR.parent / "output" / "csv"           # pasta dos CSVs
TEMPLATE_PATH = BASE_DIR / "templates" / "template.docx"  # template DOCX
OUTPUT_DIR = BASE_DIR.parent / "output"                 # pasta de saída
OUTPUT_DIR.mkdir(exist_ok=True)

# Configura localização para formatação de números (português brasileiro)
try:
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
except locale.Error:
    # Fallback para Windows
    locale.setlocale(locale.LC_ALL, 'Portuguese_Brazil.1252')

# ============================================================================
# CONFIGURAÇÃO DO SISTEMA DE LOGGING
# ============================================================================
def setup_logging():
    """
    Configura o logger para registrar eventos e erros em logs/build_report.log.
    
    Returns:
        logging.Logger: Instância configurada do logger
    """
    log_dir = BASE_DIR.parent / "logs"
    log_dir.mkdir(exist_ok=True)
    log_file = log_dir / "build_report.log"
    
    logger = logging.getLogger("build_report_logger")
    logger.setLevel(logging.INFO)
    
    file_handler = logging.FileHandler(log_file, mode='a', encoding='utf-8')
    formatter = logging.Formatter(
        '%(asctime)s - %(levelname)s - %(message)s', 
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    file_handler.setFormatter(formatter)
    
    if not logger.handlers:
        logger.addHandler(file_handler)
    
    return logger

logger = setup_logging()

# ============================================================================
# FUNÇÕES PÚBLICAS (INTERFACE DO MÓDULO)
# ============================================================================

def get_tipos_relatorio() -> list:
    """
    Retorna a lista de tipos de relatório disponíveis para o menu.
    
    Returns:
        list[str]: Lista com os tipos de relatório suportados
    
    Example:
        >>> tipos = get_tipos_relatorio()
        >>> print(tipos)
        ['Nacional', 'Regional', 'Por Setor', 'Por Grupo', 'Por Marca', 'Por PDV']
    """
    return ["Nacional", "Regional", "Por Setor", "Por Grupo", "Por Marca", "Por PDV"]

def get_opcoes_especificas(tipo: str, dados: dict) -> list:
    """
    Busca e retorna as opções de filtro para um tipo de relatório específico.
    
    Por exemplo, se o tipo for "Regional", retorna todas as regiões disponíveis
    nos dados (SUL, SUDESTE, etc.).
    
    Args:
        tipo (str): Tipo do relatório ('Regional', 'Por PDV', etc.)
        dados (dict): Dicionário com DataFrames carregados
    
    Returns:
        list: Lista ordenada de opções disponíveis
    
    Raises:
        ValueError: Se o DataFrame correspondente não for encontrado
    
    Example:
        >>> opcoes = get_opcoes_especificas('Regional', dados)
        >>> print(opcoes)
        ['CENTRO-OESTE', 'NORDESTE', 'NORTE', 'SUDESTE', 'SUL']
    """
    # Converte tipo para nome do DataFrame
    # Ex: "Por Setor" -> "visao_setor"
    nome_df = f"visao_{tipo.lower().replace('por ', '')}"
    
    df_opcoes = dados.get(nome_df)
    if df_opcoes is None:
        raise ValueError(f"DataFrame '{nome_df}.csv' não encontrado para buscar opções.")
    
    # Primeira coluna sempre é a coluna de identificação
    coluna_opcoes = list(df_opcoes.columns)[0]
    
    # Retorna lista única e ordenada
    return sorted(df_opcoes[coluna_opcoes].unique())

def carregar_dados() -> dict[str, pd.DataFrame]:
    """
    Lê todos os arquivos CSV da pasta de output e os carrega em DataFrames.
    
    Além de carregar os CSVs, esta função também:
    - Converte colunas numéricas explicitamente para int
    - Trata valores ausentes (NaN) como 0
    - Indexa os DataFrames por nome de arquivo
    
    Returns:
        dict[str, pd.DataFrame]: Dicionário com DataFrames indexados por nome
            Chaves: 'visao_vendedor', 'visao_pdv', etc.
            Valores: DataFrames correspondentes
    
    Raises:
        FileNotFoundError: Se o diretório de dados não existir ou estiver vazio
    
    Example:
        >>> dados = carregar_dados()
        >>> print(dados.keys())
        dict_keys(['visao_vendedor', 'visao_pdv', 'visao_regional', ...])
        >>> print(len(dados['visao_vendedor']))
        150
    """
    logger.info("Iniciando carregamento dos dados CSV.")
    
    dataframes = {}
    
    # Valida se o diretório existe e contém arquivos
    if not DATA_DIR.exists() or not any(DATA_DIR.iterdir()):
        raise FileNotFoundError(
            f"Diretório de dados '{DATA_DIR}' não encontrado. Execute o ETL primeiro."
        )
    
    # Define colunas que devem ser convertidas para numérico
    colunas_numericas = [
        'qtd_vendedores', 'qtd_leads', 'leads_visualizado', 'convite_enviado', 
        'convite_pendente_confirmacao', 'convite_declinado_confirmacao', 
        'convite_confirmado', 'presenca', 'testdrive', 'venda'
    ]
    
    # Carrega cada arquivo CSV
    for csv_file in DATA_DIR.glob("*.csv"):
        nome_visao = csv_file.stem  # Nome sem extensão
        df = pd.read_csv(csv_file)
        
        # Converte colunas numéricas explicitamente
        for col in colunas_numericas:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0).astype(int)
        
        dataframes[nome_visao] = df
        logger.info(f"DataFrame '{nome_visao}' carregado com {len(df)} linhas.")
    
    logger.info("Carregamento de dados concluído.")
    return dataframes

def gerar_relatorio(tipo: str, dados: dict, contexto_evento: dict, 
                    valor_filtro: str = None, data_inicio_br: str = None, 
                    data_fim_br: str = None) -> Path:
    """
    Orquestra a criação do relatório completo (DOCX e PDF).
    
    Esta é a função principal do módulo, que coordena todo o processo:
    1. Prepara contexto e seleciona dados
    2. Carrega e preenche template DOCX
    3. Popula tabelas com dados
    4. Salva arquivo DOCX
    5. Converte para PDF
    
    Args:
        tipo (str): Tipo do relatório ('Nacional', 'Regional', etc.)
        dados (dict): Dicionário com DataFrames carregados
        contexto_evento (dict): Informações do evento selecionado
            Deve conter: 'event_name', 'db_name', 'start_date', 'end_date'
        valor_filtro (str, optional): Filtro específico (ex: 'SUDESTE')
        data_inicio_br (str, optional): Data inicial no formato dd/mm/aaaa
        data_fim_br (str, optional): Data final no formato dd/mm/aaaa
    
    Returns:
        Path: Caminho do arquivo PDF gerado
    
    Example:
        >>> caminho = gerar_relatorio(
        ...     tipo='Regional',
        ...     dados=dados_carregados,
        ...     contexto_evento={'event_name': 'Liga RAM', ...},
        ...     valor_filtro='SUDESTE',
        ...     data_inicio_br='01/08/2024',
        ...     data_fim_br='31/08/2024'
        ... )
        >>> print(caminho)
        PosixPath('output/Relatorio_Regional_SUDESTE_202410141530.pdf')
    """
    logger.info("="*50)
    logger.info("MÓDULO DE GERAÇÃO DE RELATÓRIO INICIADO.")
    
    # ========================================================================
    # FASE 1: PREPARAÇÃO DE DADOS E CONTEXTO
    # ========================================================================
    contexto, df_tabela_geral = _preparar_contexto_relatorio(
        tipo, dados, contexto_evento, valor_filtro, data_inicio_br, data_fim_br
    )
    
    # ========================================================================
    # FASE 2: PREENCHIMENTO DO TEMPLATE
    # ========================================================================
    # Carrega o template DOCX
    doc = docx.Document(TEMPLATE_PATH)
    
    # Substitui todos os placeholders no documento
    replace_text_in_doc(doc, contexto)
    
    # Popula a tabela geral (única tabela - índice 0)
    if df_tabela_geral is not None:
        popular_tabela_geral(doc, df_tabela_geral, 0)

    # ========================================================================
    # FASE 3: SALVAMENTO E CONVERSÃO
    # ========================================================================
    # Define nome do arquivo baseado no tipo e timestamp
    titulo_visao = contexto["{{tipo_visao}}"]
    nome_arquivo_base = f"Relatorio_{titulo_visao}_{datetime.now().strftime('%Y%m%d%H%M')}"
    caminho_docx = OUTPUT_DIR / f"{nome_arquivo_base}.docx"
    caminho_pdf = OUTPUT_DIR / f"{nome_arquivo_base}.pdf"

    # Salva documento DOCX
    logger.info(f"Salvando relatório em DOCX: {caminho_docx}")
    doc.save(caminho_docx)
    
    # Converte para PDF
    logger.info(f"Convertendo para PDF: {caminho_pdf}")
    docx2pdf.convert(str(caminho_docx), str(caminho_pdf))
    
    logger.info("Relatório gerado com sucesso!")
    print("\nRelatório gerado com sucesso!")
    
    return caminho_pdf

# ============================================================================
# FUNÇÕES AUXILIARES INTERNAS (PRIVADAS)
# ============================================================================

def replace_text_in_doc(doc, replacements: dict):
    """
    Substitui placeholders (ex: {{chave}}) em parágrafos e tabelas do documento.
    
    Percorre todo o documento Word procurando por placeholders no formato
    {{nome_variavel}} e substitui pelos valores correspondentes.
    
    Args:
        doc (docx.Document): Documento Word a ser modificado
        replacements (dict): Dicionário de substituições
            Chave: placeholder (ex: '{{total_contatos}}')
            Valor: texto de substituição (ex: '1.234')
    """
    # Substitui em parágrafos
    for p in doc.paragraphs:
        # Reconstrói o texto completo do parágrafo (necessário pois pode estar fragmentado em runs)
        full_text = "".join(run.text for run in p.runs)
        
        # Verifica se algum placeholder está presente
        if any(key in full_text for key in replacements):
            # Substitui todos os placeholders
            for key, value in replacements.items():
                if key in full_text:
                    full_text = full_text.replace(key, str(value))
            
            # Limpa o parágrafo e reescreve com texto substituído
            p.clear()
            p.add_run(full_text)
    
    # Substitui em tabelas (caso existam placeholders em células)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Recursivamente substitui texto nas células
                replace_text_in_doc(cell, replacements)

def popular_tabela_geral(doc, df_dados: pd.DataFrame, table_index: int):
    """
    Adiciona TODAS as linhas com TODAS as métricas na tabela geral.
    
    Esta função:
    - Insere todas as linhas do DataFrame na tabela
    - Calcula colunas derivadas (Pendentes, Sem Resposta)
    - Aplica formatação profissional (Arial, tamanho 9, centralizado)
    - Garante ordem correta das colunas
    
    Args:
        doc (docx.Document): Documento Word
        df_dados (pd.DataFrame): DataFrame com os dados a serem inseridos
        table_index (int): Índice da tabela no documento (0 para primeira tabela)
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Validações iniciais
    if df_dados is None or df_dados.empty:
        logger.warning(f"Nenhum dado disponível para popular a tabela geral (índice {table_index}).")
        return
    
    if len(doc.tables) <= table_index:
        logger.error(f"Erro: O template não contém a tabela de índice {table_index}.")
        return
    
    tabela = doc.tables[table_index]
    
    # Define a coluna de identificação (primeira coluna do DataFrame)
    coluna_id = df_dados.columns[0]
    
    logger.info(f"=== POPULAR_TABELA_GERAL ===")
    logger.info(f"Coluna ID: {coluna_id}")
    logger.info(f"Total de linhas: {len(df_dados)}")
    
    # Cria cópia do DataFrame para não modificar o original
    df_tabela = df_dados.copy()
    
    # ========================================================================
    # CÁLCULO DE COLUNAS DERIVADAS
    # ========================================================================
    df_tabela['pendentes'] = df_tabela['qtd_leads'] - df_tabela['convite_enviado']
    df_tabela['sem_resposta'] = (df_tabela['convite_enviado'] - 
                                   df_tabela['convite_confirmado'] - 
                                   df_tabela['convite_declinado_confirmacao'])
    
    # ========================================================================
    # DEFINIÇÃO DA ORDEM DAS COLUNAS
    # ========================================================================
    colunas_ordenadas = [
        coluna_id,                          # Coluna identificadora
        'qtd_vendedores',                   # Métricas principais
        'qtd_leads',
        'convite_enviado',
        'pendentes',                        # Colunas derivadas
        'convite_confirmado',
        'convite_declinado_confirmacao',
        'sem_resposta'
    ]
    
    # Valida se todas as colunas existem
    colunas_faltando = [col for col in colunas_ordenadas if col not in df_tabela.columns]
    if colunas_faltando:
        logger.error(f"ERRO: Colunas faltando: {colunas_faltando}")
        return
    
    # Seleciona apenas as colunas necessárias
    df_tabela = df_tabela[colunas_ordenadas]
    
    # Ordena por nome da primeira coluna
    df_tabela = df_tabela.sort_values(by=coluna_id)
    
    logger.info(f"Primeira linha: {df_tabela.iloc[0].to_dict()}")
    
    # ========================================================================
    # PREENCHIMENTO DA TABELA COM FORMATAÇÃO
    # ========================================================================
    for _, row_data in df_tabela.iterrows():
        # Adiciona nova linha na tabela
        nova_linha = tabela.add_row().cells
        
        # Preenche cada célula da linha
        for i, coluna in enumerate(colunas_ordenadas):
            valor = row_data[coluna]
            
            # Converte para string, formatando números como inteiros
            if pd.isna(valor):
                texto = "0"
            elif isinstance(valor, (int, float)):
                texto = str(int(valor))
            else:
                texto = str(valor)
            
            # Define o texto na célula
            nova_linha[i].text = texto
            
            # Aplica formatação ao parágrafo da célula
            if len(nova_linha[i].paragraphs) > 0:
                paragrafo = nova_linha[i].paragraphs[0]
                
                # Centraliza o texto
                paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Aplica formatação à fonte (Arial, tamanho 9)
                for run in paragrafo.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
    
    logger.info(f"Tabela geral: {len(df_tabela)} linhas inseridas com formatação!")
    logger.info(f"=== FIM POPULAR_TABELA_GERAL ===")

def safe_division(numerator, denominator):
    """
    Realiza divisão de forma segura, retornando 0.0 se o denominador for zero.
    
    Esta função evita erros de ZeroDivisionError ao calcular percentuais,
    retornando 0.0 quando não é possível dividir.
    
    Args:
        numerator (float): Numerador da divisão
        denominator (float): Denominador da divisão
    
    Returns:
        float: Resultado da divisão ou 0.0 se denominador for zero
    
    Example:
        >>> safe_division(10, 2)
        5.0
        >>> safe_division(10, 0)
        0.0
    """
    return numerator / denominator if denominator != 0 else 0.0

def _preparar_contexto_relatorio(tipo: str, dados: dict, contexto_evento: dict, 
                                  valor_filtro: str = None, data_inicio_br: str = None, 
                                  data_fim_br: str = None):
    """
    Prepara todos os dados e cálculos necessários para o relatório.
    
    Esta é a função mais complexa do módulo. Ela implementa a lógica de:
    - Seleção dos dados corretos baseado no tipo e filtro
    - Determinação de qual DataFrame usar para a tabela geral
    - Cálculo de todas as métricas e percentuais
    - Preparação do contexto de substituição de placeholders
    
    Args:
        tipo (str): Tipo do relatório
        dados (dict): Dicionário com DataFrames
        contexto_evento (dict): Informações do evento
        valor_filtro (str, optional): Filtro específico
        data_inicio_br (str, optional): Data inicial (dd/mm/aaaa)
        data_fim_br (str, optional): Data final (dd/mm/aaaa)
    
    Returns:
        tuple: (contexto, df_tabela_geral)
            - contexto (dict): Dicionário com todos os placeholders e valores
            - df_tabela_geral (pd.DataFrame): DataFrame para a tabela detalhada
    """
    # ========================================================================
    # SELEÇÃO DE DADOS PRINCIPAIS
    # ========================================================================
    dados_relatorio = None
    titulo_visao = tipo
    
    if valor_filtro:
        # Relatório específico (ex: Regional SUDESTE)
        nome_df_base = f"visao_{tipo.lower().replace('por ', '')}"
        df_base = dados.get(nome_df_base)
        coluna_filtro = list(df_base.columns)[0]
        dados_relatorio = df_base[df_base[coluna_filtro] == valor_filtro].iloc[0]
        titulo_visao = f"{tipo.replace('Por ','')} - {valor_filtro}"
    else: 
        # Relatório acumulado (ex: Regional Acumulado)
        df_nacional = dados.get("visao_nacional")
        dados_relatorio = df_nacional.iloc[0]
        titulo_visao = f"{tipo.replace('Por ','')} (Acumulado)"
        if tipo == "Nacional": 
            titulo_visao = "Nacional"

    if dados_relatorio is None:
        raise ValueError(
            f"Não foram encontrados dados para a seleção: Tipo='{tipo}', Filtro='{valor_filtro}'"
        )

    # ========================================================================
    # SELEÇÃO DE DADOS PARA TABELA GERAL
    # ========================================================================
    df_tabela_geral = None
    nome_coluna_tabela_geral = "Categoria"
    
    if tipo == 'Nacional':
        # Tabela Geral: TODAS as regiões
        df_tabela_geral = dados.get("visao_regional")
        nome_coluna_tabela_geral = "Região"
        
    elif tipo == 'Regional':
        if valor_filtro:
            # Regional Específico - TODOS os setores da região
            df_setor = dados.get("visao_setor")
            df_pdv_da_regiao = dados.get("visao_pdv")[dados.get("visao_pdv")['rid'] == valor_filtro]
            setores_da_regiao = df_pdv_da_regiao['sid'].unique()
            df_tabela_geral = df_setor[df_setor['sid'].isin(setores_da_regiao)]
            nome_coluna_tabela_geral = "Setor"
        else:
            # Regional Acumulado - TODOS os setores
            df_tabela_geral = dados.get("visao_setor")
            nome_coluna_tabela_geral = "Setor"
            
    elif tipo in ['Por Setor', 'Por Grupo', 'Por Marca']:
        df_pdv = dados.get("visao_pdv")
        nome_df_base = f"visao_{tipo.lower().replace('por ', '')}"
        coluna_filtro_pa = list(dados.get(nome_df_base).columns)[0]
        
        if valor_filtro:
            # Filtro específico - TODOS os PDVs do filtro
            df_tabela_geral = df_pdv[df_pdv[coluna_filtro_pa] == valor_filtro]
        else:
            # Acumulado - TODOS os PDVs
            df_tabela_geral = df_pdv
            
        nome_coluna_tabela_geral = "PDV"
        
    elif tipo == 'Por PDV':
        # TODOS os vendedores do PDV específico
        df_vendedor = dados.get("visao_vendedor")
        df_tabela_geral = df_vendedor[df_vendedor['pdv'] == valor_filtro]
        nome_coluna_tabela_geral = "Vendedor"
        
    # ========================================================================
    # CÁLCULO DE MÉTRICAS
    # ========================================================================
    total_enviados = dados_relatorio['convite_enviado']
    total_confirmados = dados_relatorio['convite_confirmado']
    total_presencas = dados_relatorio['presenca']
    total_testdrives = dados_relatorio['testdrive']
    total_vendas = dados_relatorio['venda']

    # ========================================================================
    # PREPARAÇÃO DO CONTEXTO DE SUBSTITUIÇÃO
    # ========================================================================
    contexto = {
        "{{nome_evento}}": contexto_evento['event_name'],
        "{{data_inicio}}": data_inicio_br if data_inicio_br else "N/A",
        "{{data_fim}}": data_fim_br if data_fim_br else "N/A",
        "{{tipo_visao}}": titulo_visao,
        "{{categoria_visao}}": nome_coluna_tabela_geral,
        
        # Métricas absolutas (formatadas com separador de milhares)
        "{{total_contatos}}": locale.format_string("%d", dados_relatorio['qtd_leads'], grouping=True),
        "{{total_nao_vistos}}": locale.format_string("%d", dados_relatorio['qtd_leads'] - dados_relatorio['leads_visualizado'], grouping=True),
        "{{total_enviados}}": locale.format_string("%d", total_enviados, grouping=True),
        "{{total_envio_pendente}}": locale.format_string("%d", dados_relatorio['qtd_leads'] - total_enviados, grouping=True),
        "{{total_confirmados}}": locale.format_string("%d", total_confirmados, grouping=True),
        "{{total_declinados}}": locale.format_string("%d", dados_relatorio['convite_declinado_confirmacao'], grouping=True),
        "{{total_sem_resposta}}": locale.format_string("%d", total_enviados - total_confirmados - dados_relatorio['convite_declinado_confirmacao'], grouping=True),
        "{{total_vendedores}}": locale.format_string("%d", dados_relatorio.get('qtd_vendedores', 0), grouping=True),
        "{{total_presencas}}": locale.format_string("%d", total_presencas, grouping=True),
        "{{total_testdrives}}": locale.format_string("%d", total_testdrives, grouping=True),
        "{{total_vendas}}": locale.format_string("%d", total_vendas, grouping=True),
        
        # Percentuais de conversão (formatados com vírgula decimal)
        "{{perc_enviados_confirmados}}": f"{safe_division(total_confirmados, total_enviados) * 100:.2f}".replace('.',','),
        "{{perc_confirmados_presencas}}": f"{safe_division(total_presencas, total_confirmados) * 100:.2f}".replace('.',','),
        "{{perc_presencas_testdrives}}": f"{safe_division(total_testdrives, total_presencas) * 100:.2f}".replace('.',','),
        "{{perc_presencas_vendas}}": f"{safe_division(total_vendas, total_presencas) * 100:.2f}".replace('.',','),
        "{{perc_testdrives_vendas}}": f"{safe_division(total_vendas, total_testdrives) * 100:.2f}".replace('.',','),
    }
    
    return contexto, df_tabela_geral

# ============================================================================
# PONTO DE ENTRADA PARA TESTES ISOLADOS
# ============================================================================
if __name__ == "__main__":
    """
    Este bloco só é executado se o script for chamado diretamente.
    """
    print("Este é um módulo de serviço do Gerador de Relatórios.")
    print("Execute o `main.py` para iniciar o pipeline completo.")
    pass